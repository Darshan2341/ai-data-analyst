import streamlit as st
import pandas as pd
import plotly.express as px
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage
from dotenv import load_dotenv
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt
import tempfile
import os

load_dotenv()

st.set_page_config(page_title="AI Data Analyst", page_icon="📊", layout="wide")

st.markdown("""
<style>
    .main { background-color: #0f1117; }
    .block-container { padding: 2rem 3rem; }
    .insight-card {
        background-color: #1a1d27;
        border-radius: 10px;
        padding: 1.2rem;
        margin-bottom: 0.8rem;
        border-left: 4px solid #4fc3f7;
        color: #fff;
    }
    .stat-card {
        background-color: #1a1d27;
        border-radius: 10px;
        padding: 1rem;
        text-align: center;
        border: 1px solid #2a3550;
    }
    .stat-number { font-size: 1.8rem; font-weight: 700; color: #4fc3f7; }
    .stat-label { color: #aaa; font-size: 0.85rem; }
</style>
""", unsafe_allow_html=True)


# ── Helper functions ──────────────────────────────────────────

def get_df_summary(df):
    return f"""
Dataset shape: {df.shape[0]} rows, {df.shape[1]} columns
Columns: {', '.join(df.columns.tolist())}
Column types:\n{df.dtypes.to_string()}
Basic statistics:\n{df.describe().to_string()}
First 5 rows:\n{df.head().to_string()}
Missing values:\n{df.isnull().sum().to_string()}
"""

def ask_ai(df, question):
    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0)
    prompt = f"""You are an expert data analyst. Answer clearly based on the data.
DATASET:\n{get_df_summary(df)}
QUESTION: {question}
ANSWER:"""
    return llm.invoke([HumanMessage(content=prompt)]).content

def get_ai_insights(df):
    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0)
    prompt = f"""You are an expert data analyst. Give exactly 5 key insights from this dataset.
Each insight should be specific with actual numbers.
Format as:
1. [insight]
2. [insight]
3. [insight]
4. [insight]
5. [insight]

DATASET:\n{get_df_summary(df)}

5 KEY INSIGHTS:"""
    return llm.invoke([HumanMessage(content=prompt)]).content

def auto_charts(df):
    charts = []
    numeric_cols = df.select_dtypes(include='number').columns.tolist()
    categorical_cols = df.select_dtypes(include='object').columns.tolist()
    date_cols = [c for c in df.columns if 'date' in c.lower() or 'time' in c.lower()]

    if categorical_cols and numeric_cols:
        top = df.groupby(categorical_cols[0])[numeric_cols[0]].sum().nlargest(15).reset_index()
        charts.append(("Bar Chart", px.bar(top, x=categorical_cols[0], y=numeric_cols[0],
            title=f"{numeric_cols[0]} by {categorical_cols[0]}", color=numeric_cols[0],
            color_continuous_scale="Blues", template="plotly_dark")))

    if date_cols and numeric_cols:
        try:
            df[date_cols[0]] = pd.to_datetime(df[date_cols[0]])
            line_data = df.groupby(date_cols[0])[numeric_cols[0]].sum().reset_index()
            charts.append(("Trend Over Time", px.line(line_data, x=date_cols[0], y=numeric_cols[0],
                title=f"{numeric_cols[0]} over time", template="plotly_dark")))
        except: pass

    if categorical_cols and numeric_cols:
        pie_data = df.groupby(categorical_cols[0])[numeric_cols[0]].sum().nlargest(8).reset_index()
        charts.append(("Distribution", px.pie(pie_data, names=categorical_cols[0], values=numeric_cols[0],
            title=f"{numeric_cols[0]} Distribution", template="plotly_dark")))

    if numeric_cols:
        charts.append(("Histogram", px.histogram(df, x=numeric_cols[0],
            title=f"Distribution of {numeric_cols[0]}", template="plotly_dark",
            color_discrete_sequence=["#4fc3f7"])))

    if len(numeric_cols) >= 2:
        charts.append(("Scatter", px.scatter(df, x=numeric_cols[0], y=numeric_cols[1],
            title=f"{numeric_cols[0]} vs {numeric_cols[1]}", template="plotly_dark",
            color_discrete_sequence=["#4fc3f7"])))
    return charts


# ── Export functions ──────────────────────────────────────────

def export_excel_sheets(df, insights, filename):
    """Export as Excel with multiple formatted sheets."""
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:

        # Sheet 1 - Raw Data
        df.to_excel(writer, sheet_name='Raw Data', index=False)
        ws = writer.sheets['Raw Data']
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E79")
            cell.alignment = Alignment(horizontal="center")

        # Sheet 2 - Statistics
        df.describe().reset_index().to_excel(writer, sheet_name='Statistics', index=False)
        ws2 = writer.sheets['Statistics']
        for cell in ws2[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E79")

        # Sheet 3 - Summary by category
        cat_cols = df.select_dtypes(include='object').columns.tolist()
        num_cols = df.select_dtypes(include='number').columns.tolist()
        if cat_cols and num_cols:
            summary = df.groupby(cat_cols[0])[num_cols].sum().reset_index()
            summary.to_excel(writer, sheet_name='Summary', index=False)
            ws3 = writer.sheets['Summary']
            for cell in ws3[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="1F4E79")

        # Sheet 4 - AI Insights
        lines = [l.strip() for l in insights.split('\n') if l.strip()]
        insights_df = pd.DataFrame({'AI Insights': lines})
        insights_df.to_excel(writer, sheet_name='AI Insights', index=False)
        ws4 = writer.sheets['AI Insights']
        for cell in ws4[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E79")

    return output.getvalue()


def export_word_doc(df, insights, filename):
    """Export as a clean Word document report."""
    doc = Document()

    # Title
    title = doc.add_heading('Data Analysis Report', 0)
    title.alignment = 1  # center

    doc.add_paragraph(f'File: {filename}')
    doc.add_paragraph(f'Rows: {df.shape[0]}  |  Columns: {df.shape[1]}  |  Missing Values: {df.isnull().sum().sum()}')
    doc.add_paragraph('')

    # Dataset overview
    doc.add_heading('Dataset Overview', level=1)
    doc.add_paragraph(f'Columns: {", ".join(df.columns.tolist())}')

    # Statistics table
    doc.add_heading('Summary Statistics', level=1)
    stats = df.describe().reset_index()
    table = doc.add_table(rows=1, cols=len(stats.columns))
    table.style = 'Table Grid'
    for i, col in enumerate(stats.columns):
        table.rows[0].cells[i].text = str(col)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for _, row in stats.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, float) else str(val)

    doc.add_paragraph('')

    # AI Insights
    doc.add_heading('AI-Generated Insights', level=1)
    lines = [l.strip() for l in insights.split('\n') if l.strip() and l.strip()[0].isdigit()]
    for line in lines:
        doc.add_paragraph(line, style='List Number')

    # Raw data sample
    doc.add_heading('Data Sample (First 10 Rows)', level=1)
    sample = df.head(10).reset_index(drop=True)
    table2 = doc.add_table(rows=1, cols=len(sample.columns))
    table2.style = 'Table Grid'
    for i, col in enumerate(sample.columns):
        table2.rows[0].cells[i].text = str(col)
        table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for _, row in sample.iterrows():
        row_cells = table2.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Save to bytes
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ── Main UI ───────────────────────────────────────────────────

st.title("📊 AI Data Analysis Assistant")
st.markdown("Upload any CSV — get instant charts, AI insights, and export to **Excel or Word**.")

uploaded_file = st.file_uploader("Upload your CSV file", type=["csv"])

if uploaded_file:
    df = pd.read_csv(uploaded_file, sep=";",engine="python")
    numeric_cols = df.select_dtypes(include='number').columns.tolist()

    st.markdown("---")

    # Quick stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{df.shape[0]:,}</div><div class="stat-label">Total Rows</div></div>', unsafe_allow_html=True)
    with col2:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{df.shape[1]}</div><div class="stat-label">Columns</div></div>', unsafe_allow_html=True)
    with col3:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{df.isnull().sum().sum()}</div><div class="stat-label">Missing Values</div></div>', unsafe_allow_html=True)
    with col4:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{len(numeric_cols)}</div><div class="stat-label">Numeric Columns</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Data", "📈 Charts", "🤖 AI Insights", "💬 Ask AI", "📥 Export"])

    with tab1:
        st.subheader("Raw Data")
        st.dataframe(df, use_container_width=True)
        st.subheader("Summary Statistics")
        st.dataframe(df.describe(), use_container_width=True)

    with tab2:
        st.subheader("Auto-Generated Charts")
        charts = auto_charts(df)
        for i in range(0, len(charts), 2):
            cols = st.columns(2)
            for j, col in enumerate(cols):
                if i + j < len(charts):
                    name, fig = charts[i + j]
                    with col:
                        st.plotly_chart(fig, use_container_width=True)

    with tab3:
        st.subheader("AI-Generated Insights")
        if st.button("Generate Insights", use_container_width=True):
            with st.spinner("AI is analyzing your data..."):
                insights = get_ai_insights(df)
            st.session_state.insights = insights
            lines = [l.strip() for l in insights.split('\n') if l.strip() and l.strip()[0].isdigit()]
            for line in lines:
                st.markdown(f'<div class="insight-card">{line}</div>', unsafe_allow_html=True)

    with tab4:
        st.subheader("Ask anything about your data")
        question = st.text_input("Your question:", placeholder="e.g. Which product has the highest sales?")
        if st.button("Ask AI", use_container_width=True):
            if question:
                with st.spinner("Thinking..."):
                    answer = ask_ai(df, question)
                st.markdown(f'<div class="insight-card">💬 {answer}</div>', unsafe_allow_html=True)

    with tab5:
        st.subheader("Export Your Analysis")
        st.markdown("Choose the format you want to download:")

        # Ask user what format they want
        export_format = st.radio(
            "What format do you need the output in?",
            ["📊 Excel Spreadsheet (multiple sheets)", "📝 Word Document (blank doc style report)"],
            index=0
        )

        export_name = st.text_input("File name (optional):", value=uploaded_file.name.replace('.csv', ''))

        if st.button("Generate & Download", use_container_width=True):
            with st.spinner("Generating your report... this takes about 15 seconds"):
                insights = get_ai_insights(df)

            if "Excel" in export_format:
                excel_data = export_excel_sheets(df, insights, uploaded_file.name)
                st.download_button(
                    "⬇️ Download Excel File (.xlsx)",
                    data=excel_data,
                    file_name=f"{export_name}_analysis.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
                st.success("Excel file ready! It has 4 sheets: Raw Data, Statistics, Summary, AI Insights.")

            elif "Word" in export_format:
                word_data = export_word_doc(df, insights, uploaded_file.name)
                st.download_button(
                    "⬇️ Download Word Document (.docx)",
                    data=word_data,
                    file_name=f"{export_name}_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("Word document ready! It has statistics table, AI insights, and data sample.")

else:
    st.info("Upload a CSV file to get started!")
    st.markdown("### Don't have a CSV? Download a free dataset from:")
    st.markdown("- [Kaggle Datasets](https://www.kaggle.com/datasets)")
    st.markdown("- [UCI ML Repository](https://archive.ics.uci.edu/)")
